import os
import json
import logging
import re
import subprocess
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Optional, List, Tuple, Any
from dataclasses import dataclass, asdict
import base64

# --- Dependency Imports with Graceful Fallbacks ---
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import ollama
except ImportError:
    ollama = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

import io
try:
    from PIL import Image
except ImportError:
    Image = None


logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

from .memory import DocumentTableMemory, TableFingerprint
from .ocr_engine import OllamaOCREngine

class IFileExtractor(ABC):
    """Abstract interface for all file extraction strategies."""
    @abstractmethod
    def extract(self, file_path: Path) -> str:
        pass


class TxtExtractor(IFileExtractor):
    def extract(self, file_path: Path) -> str:
        try:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            safe_content = content.replace("<", "&lt;").replace(">", "&gt;")
            return f"<!DOCTYPE html><html><body><h1>{file_path.name}</h1><pre>{safe_content}</pre></body></html>"
        except Exception as e:
            return f"<p>Error: {e}</p>"


class DocxExtractor(IFileExtractor):
    def extract(self, file_path: Path) -> str:
        if not Document:
            return "<p>Error: 'python-docx' library missing.</p>"
        try:
            doc = Document(file_path)
            html = [f"<!DOCTYPE html><html><body><h1>{file_path.name}</h1>"]
            for para in doc.paragraphs:
                if para.text.strip():
                    html.append(f"<p>{para.text}</p>")
            for table in doc.tables:
                html.append("<table border='1' style='border-collapse: collapse; width: 100%;'>")
                for row in table.rows:
                    html.append("<tr>")
                    for cell in row.cells:
                        html.append(f"<td style='border: 1px solid black; padding: 10px;'>{cell.text}</td>")
                    html.append("</tr>")
                html.append("</table>")
            html.append("</body></html>")
            return "\n".join(html)
        except Exception as e:
            return f"<p>Error: {e}</p>"


class ImageExtractor(IFileExtractor):
    def __init__(self, ocr_engine: OllamaOCREngine):
        self.ocr_engine = ocr_engine

    def extract(self, file_path: Path) -> str:
        if not Image: return "<p>Error: PIL missing.</p>"
        try:
            img_bytes = file_path.read_bytes()
            img = Image.open(io.BytesIO(img_bytes))

            w, h = img.size
            if h > w * 1.2:
                content = self.ocr_engine.perform_regional_ocr(img)
            else:
                content = self.ocr_engine.perform_ocr(img_bytes)

            return f"<!DOCTYPE html><html><body><h1>{file_path.name}</h1>{content}</body></html>"
        except Exception as e:
            return f"<p>Error: {e}</p>"


class PDFExtractorV4(IFileExtractor):
    def __init__(self, ocr_engine: OllamaOCREngine):
        self.ocr_engine = ocr_engine

    def _is_scanned(self, doc) -> bool:
        check_pages = min(3, len(doc))
        total_text_len = sum(len(doc[i].get_text().strip()) for i in range(check_pages))
        avg_text_len = total_text_len / check_pages if check_pages > 0 else 0
        return avg_text_len < 500

    def _convert_table_to_html(self, table, header_override: Optional[str] = None) -> Tuple[str, str]:
        try:
            processed_cells = set()
            extracted = table.extract()
            body_html = []
            header_html = []

            local_header_exists = (table.header and any(c for c in table.header.cells if c))

            if header_override:
                header_html.append(header_override)
                data_start_idx = 1 if local_header_exists else 0
            else:
                if local_header_exists:
                    header_html.append("<thead><tr>")
                    for i, cell in enumerate(table.header.cells):
                        if not cell: continue

                        vspan = getattr(cell, "vspan", (0, 1))
                        hspan = getattr(cell, "hspan", (i, i+1))

                        cell_id = (vspan[0], hspan[0])
                        if cell_id in processed_cells: continue
                        processed_cells.add(cell_id)

                        rowspan = vspan[1] - vspan[0]
                        colspan = hspan[1] - hspan[0]

                        attrs = ""
                        if rowspan > 1: attrs += f" rowspan='{rowspan}'"
                        if colspan > 1: attrs += f" colspan='{colspan}'"
                        content = extracted[0][i] if extracted and i < len(extracted[0]) else ""
                        header_html.append(f"<th style='background-color: #f2f2f2; padding: 8px;'{attrs}>{content}</th>")
                    header_html.append("</tr></thead>")
                    data_start_idx = 1
                else:
                    data_start_idx = 0

            body_html.append("<tbody>")
            for r_idx in range(data_start_idx, len(extracted)):
                row_raw = extracted[r_idx]
                row_obj_idx = r_idx - 1 if local_header_exists else r_idx
                row_obj = table.rows[row_obj_idx] if 0 <= row_obj_idx < len(table.rows) else None

                body_html.append("<tr>")
                if row_obj:
                    for c_idx, cell in enumerate(row_obj.cells):
                        if not cell: continue

                        vspan = getattr(cell, "vspan", (r_idx, r_idx+1))
                        hspan = getattr(cell, "hspan", (c_idx, c_idx+1))

                        cell_id = (vspan[0], hspan[0])
                        if cell_id in processed_cells: continue
                        processed_cells.add(cell_id)

                        rowspan = vspan[1] - vspan[0]
                        colspan = hspan[1] - hspan[0]

                        attrs = ""
                        if rowspan > 1: attrs += f" rowspan='{rowspan}'"
                        if colspan > 1: attrs += f" colspan='{colspan}'"
                        content = row_raw[c_idx] if c_idx < len(row_raw) and row_raw[c_idx] else ""
                        body_html.append(f"<td style='padding: 8px;'{attrs}>{content}</td>")
                else:
                    for cell_text in row_raw:
                        body_html.append(f"<td style='padding: 8px;'>{cell_text}</td>")
                body_html.append("</tr>")

            body_html.append("</tbody>")

            final_header = "".join(header_html)
            full_html = f"<table border='1' style='border-collapse: collapse; width: 100%; margin: 20px 0;'>{final_header}{''.join(body_html)}</table>"

            return full_html, final_header
        except Exception as e:
            return "<p>Error converting table.</p>", ""

    def _extract_page_content(self, page, memory: DocumentTableMemory) -> str:
        tabs = page.find_tables()
        tables = tabs.tables
        table_bboxes = [t.bbox for t in tables]
        blocks = page.get_text("blocks")

        html = []
        elements = []
        processed_table_indices = set()

        for block in blocks:
            bx0, by0, bx1, by1, text, bno, btype = block
            clean_text = text.strip()
            if not clean_text: continue

            is_inside_table = False
            b_rect = fitz.Rect(bx0, by0, bx1, by1)

            for i, tbox in enumerate(table_bboxes):
                intersect = b_rect & fitz.Rect(tbox)
                threshold = 0.95 if by0 < 70 else 0.40
                if intersect.get_area() > (b_rect.get_area() * threshold):
                    is_inside_table = True
                    if i not in processed_table_indices:
                        elements.append(('table', tbox[1], tables[i]))
                        processed_table_indices.add(i)
                    break

            if not is_inside_table:
                elements.append(('text', by0, clean_text))

        for i, table in enumerate(tables):
            if i not in processed_table_indices:
                elements.append(('table', table.bbox[1], table))

        elements.sort(key=lambda x: x[1])

        marker_pattern = r'(\s|^)([a-z0-9]\)|[0-9]+\.[0-9]+|[0-9]+\.|[•\-\*])(\s+)'

        for el_type, y0, content in elements:
            if el_type == 'table':
                table = content
                bbox = table.bbox
                current_fingerprint = TableFingerprint(
                    col_count=table.col_count,
                    x_start=bbox[0],
                    width=bbox[2] - bbox[0],
                    header_html=""
                )

                if memory.is_continuation(current_fingerprint, page.number):
                    header_html = memory.get_header_html()
                else:
                    header_html = None

                table_html, header_only_html = self._convert_table_to_html(table, header_override=header_html)
                html.append(table_html)

                if header_html:
                    current_fingerprint.header_html = header_html
                else:
                    current_fingerprint.header_html = header_only_html

                memory.update(current_fingerprint, page.number)
            else:
                text = content.replace("\n", " ").strip()
                is_heading = len(text) < 100 and (text.isupper() or re.match(r'^[0-9]+(\.[0-9]+)*\s+[A-Z]', text))

                structured = re.sub(marker_pattern, r'</li><li><strong>\2</strong> ', text)
                if "</li><li>" in structured:
                    html.append(f"<ul>{re.sub(r'^</li><li>', '', structured)}</li></ul>")
                elif is_heading:
                    html.append(f"<h3>{text}</h3>")
                else:
                    html.append(f"<p>{text}</p>")

        return "".join(html)

    def extract(self, file_path: Path) -> str:
        if not fitz: return "<p>Error: PyMuPDF missing.</p>"
        try:
            doc = fitz.open(file_path)
            html_parts = [f"<!DOCTYPE html><html><body><h1>{file_path.name}</h1>"]
            doc_memory = DocumentTableMemory()

            if self._is_scanned(doc):
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_html = self.ocr_engine.perform_regional_ocr(img, page_num=i, memory=doc_memory)
                    html_parts.append(f"<section><h2>Page {i+1}</h2>{page_html}</section>")
            else:
                for i, page in enumerate(doc):
                    html_parts.append(self._extract_page_content(page, memory=doc_memory))

            html_parts.append("</body></html>")
            doc.close()
            return "\n".join(html_parts)
        except Exception as e:
            return f"<p>Error: {e}</p>"


class DocumentRegistry:
    def __init__(self, ocr_engine: OllamaOCREngine):
        self._registry: Dict[str, IFileExtractor] = {}
        self.ocr_engine = ocr_engine
        self._initialize_defaults()

    def _initialize_defaults(self):
        self.register('.txt', TxtExtractor())
        self.register('.docx', DocxExtractor())
        self.register('.pdf', PDFExtractorV4(self.ocr_engine))
        img_handler = ImageExtractor(self.ocr_engine)
        for ext in ['.png', '.jpg', '.jpeg']: self.register(ext, img_handler)

    def register(self, extension: str, extractor: IFileExtractor):
        self._registry[extension.lower()] = extractor

    def get_extractor(self, extension: str) -> Optional[IFileExtractor]:
        return self._registry.get(extension.lower())


